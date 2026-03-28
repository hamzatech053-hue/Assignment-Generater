"""
CV/Resume generation utilities
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def create_enhanced_cv_docx(resume_data, template_name="Professional"):
    """
    Create an enhanced CV in DOCX format with template-specific styling
    
    Args:
        resume_data: Dictionary containing resume information
        template_name: Template style to apply
        
    Returns:
        Document object or None
    """
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # Template-specific styling
    template_styles = {
        "Executive": {
            "title_size": 28,
            "title_color": RGBColor(0, 32, 96),
            "header_color": RGBColor(0, 32, 96),
            "accent_color": RGBColor(128, 128, 128)
        },
        "Professional": {
            "title_size": 24,
            "title_color": RGBColor(0, 0, 0),
            "header_color": RGBColor(0, 0, 0),
            "accent_color": RGBColor(64, 64, 64)
        },
        "Modern": {
            "title_size": 26,
            "title_color": RGBColor(0, 123, 255),
            "header_color": RGBColor(0, 123, 255),
            "accent_color": RGBColor(108, 117, 125)
        },
        "Creative Designer": {
            "title_size": 28,
            "title_color": RGBColor(220, 53, 69),
            "header_color": RGBColor(220, 53, 69),
            "accent_color": RGBColor(255, 193, 7)
        },
        "Academic Scholar": {
            "title_size": 24,
            "title_color": RGBColor(40, 167, 69),
            "header_color": RGBColor(40, 167, 69),
            "accent_color": RGBColor(108, 117, 125)
        },
        "Sales Professional": {
            "title_size": 26,
            "title_color": RGBColor(255, 193, 7),
            "header_color": RGBColor(255, 193, 7),
            "accent_color": RGBColor(0, 0, 0)
        },
        "Healthcare Expert": {
            "title_size": 24,
            "title_color": RGBColor(23, 162, 184),
            "header_color": RGBColor(23, 162, 184),
            "accent_color": RGBColor(108, 117, 125)
        }
    }
    
    style = template_styles.get(template_name, template_styles["Professional"])
    
    # Add title with template styling
    title = doc.add_paragraph()
    title_run = title.add_run(resume_data.get("name", "Resume"))
    title_run.font.size = Pt(style["title_size"])
    title_run.font.bold = True
    title_run.font.color.rgb = style["title_color"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add job title
    if "title" in resume_data:
        job_title = doc.add_paragraph()
        job_title_run = job_title.add_run(resume_data["title"])
        job_title_run.font.size = Pt(16)
        job_title_run.font.color.rgb = style["accent_color"]
        job_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add contact info
    contact_info = []
    if "email" in resume_data:
        contact_info.append(resume_data["email"])
    if "phone" in resume_data:
        contact_info.append(resume_data["phone"])
    if "location" in resume_data:
        contact_info.append(resume_data["location"])
    
    if contact_info:
        contact = doc.add_paragraph()
        contact.add_run(" | ".join(contact_info))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add sections
    if "summary" in resume_data:
        _add_section_header(doc, "PROFESSIONAL SUMMARY", style)
        doc.add_paragraph(resume_data["summary"])
    
    if "experience" in resume_data and resume_data["experience"]:
        _add_section_header(doc, "PROFESSIONAL EXPERIENCE", style)
        for exp in resume_data["experience"]:
            exp_header = doc.add_paragraph()
            exp_header.add_run(f"{exp.get('role', '')} | {exp.get('company', '')}").bold = True
            
            if exp.get('duration') or exp.get('location'):
                exp_details = doc.add_paragraph()
                details_text = []
                if exp.get('duration'):
                    details_text.append(exp['duration'])
                if exp.get('location'):
                    details_text.append(exp['location'])
                exp_details.add_run(" | ".join(details_text))
                exp_details_run = exp_details.runs[0]
                exp_details_run.font.color.rgb = style["accent_color"]
            
            if exp.get('description'):
                desc_lines = exp['description'].split('\n')
                for line in desc_lines:
                    if line.strip():
                        doc.add_paragraph(line.strip(), style='List Bullet')
    
    if "education" in resume_data and resume_data["education"]:
        _add_section_header(doc, "EDUCATION", style)
        for edu in resume_data["education"]:
            edu_para = doc.add_paragraph()
            edu_text = f"{edu.get('degree', '')} | {edu.get('institution', '')}"
            if edu.get('year'):
                edu_text += f" | {edu['year']}"
            edu_para.add_run(edu_text).bold = True
            
            if edu.get('details'):
                details_para = doc.add_paragraph(edu['details'])
                details_para.runs[0].font.color.rgb = style["accent_color"]
    
    if "skills" in resume_data and resume_data["skills"]:
        _add_section_header(doc, "CORE COMPETENCIES", style)
        skills_list = resume_data["skills"]
        if isinstance(skills_list, list):
            for i in range(0, len(skills_list), 3):
                skill_group = skills_list[i:i+3]
                doc.add_paragraph(" • ".join(skill_group))
        else:
            doc.add_paragraph(str(skills_list))
    
    if "projects" in resume_data and resume_data["projects"]:
        _add_section_header(doc, "KEY PROJECTS", style)
        for project in resume_data["projects"]:
            proj_header = doc.add_paragraph()
            proj_header.add_run(project.get('name', '')).bold = True
            
            if project.get('description'):
                doc.add_paragraph(project['description'])
            
            if project.get('achievements'):
                doc.add_paragraph(f"Results: {project['achievements']}")
    
    if "certifications" in resume_data and resume_data["certifications"]:
        _add_section_header(doc, "CERTIFICATIONS", style)
        for cert in resume_data["certifications"]:
            doc.add_paragraph(cert, style='List Bullet')
    
    return doc


def _add_section_header(doc, title, style):
    """
    Add a styled section header
    """
    header = doc.add_paragraph()
    header_run = header.add_run(title)
    header_run.font.size = Pt(14)
    header_run.font.bold = True
    header_run.font.color.rgb = style["header_color"]
    
    doc.add_paragraph()
